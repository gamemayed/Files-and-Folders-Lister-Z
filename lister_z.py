import os
import json
from docx import Document

# Language dictionaries
LANGUAGES = {
    "en": {
        "select_language": "Select language / Selecione o idioma:\n1. Português Brasileiro\n2. English\nEnter 1 or 2: ",
        "enter_directory": "Enter the directory to list: ",
        "error_directory": "Error: The directory '{dir}' does not exist. Please enter a valid folder.",
        "hide_hidden": "Do you want to hide files such as desktop.ini? (y/yes/n/no): ",
        "filter_input": "Enter sub-folder names or keywords to filter (comma-separated, case-insensitive, substring match), or leave blank to include all: ",
        "output_mode": "Do you want to generate the output as DOCX (A), TXT (B), or JSON (C)? ",
        "invalid_mode": "Invalid input. Please enter DOCX/A or TXT/B, or JSON/C.",
        "choose_option": "Choose listing option:",
        "option_1": "  1. List both folders and files",
        "option_2": "  2. List only folders",
        "option_3": "  3. List only files",
        "enter_choice": "Enter your choice (1/2/3): ",
        "invalid_choice": "Invalid choice. Please enter 1, 2, or 3.",
        "invalid_number": "Invalid input. Please enter a number (1, 2, or 3).",
        "list_success": "List generated successfully: {path}",
        "list_generated": "The List has been generated",
        "docx_failed": "DOCX generation failed: {err}. Falling back to TXT mode.",
        "json_exported": "JSON database exported: {path}",
        "credits": "Credits: User Ium101 from GitHub and AI tools",
        "press_any_button": "Press any button to exit"
    },
    "pt": {
        "select_language": "Selecione o idioma / Select language:\n1. Português Brasileiro\n2. English\nDigite 1 ou 2: ",
        "enter_directory": "Digite o diretório para listar: ",
        "error_directory": "Erro: O diretório '{dir}' não existe. Por favor, insira uma pasta válida.",
        "hide_hidden": "Deseja ocultar arquivos como desktop.ini? (s/sim/n/não): ",
        "filter_input": "Digite nomes de subpastas ou palavras-chave para filtrar (separados por vírgula, sem diferenciar maiúsculas/minúsculas), ou deixe em branco para incluir todas: ",
        "output_mode": "Deseja gerar a saída como DOCX (A), TXT (B) ou JSON (C)? ",
        "invalid_mode": "Entrada inválida. Por favor, insira DOCX/A, TXT/B ou JSON/C.",
        "choose_option": "Escolha a opção de listagem:",
        "option_1": "  1. Listar pastas e arquivos",
        "option_2": "  2. Listar apenas pastas",
        "option_3": "  3. Listar apenas arquivos",
        "enter_choice": "Digite sua escolha (1/2/3): ",
        "invalid_choice": "Escolha inválida. Por favor, insira 1, 2 ou 3.",
        "invalid_number": "Entrada inválida. Por favor, insira um número (1, 2 ou 3).",
        "list_success": "Lista gerada com sucesso: {path}",
        "list_generated": "A lista foi gerada",
        "docx_failed": "Falha ao gerar DOCX: {err}. Gerando TXT em vez disso.",
        "json_exported": "Banco de dados JSON exportado: {path}",
        "credits": "Créditos: Usuário Ium101 do GitHub e Ferramentas IA",
        "press_any_button": "Pressione qualquer botão para sair"
    }
}

def get_lang():
    while True:
        lang_choice = input(LANGUAGES["en"]["select_language"])
        if lang_choice.strip() == "1":
            return "pt"
        elif lang_choice.strip() == "2":
            return "en"
        else:
            print("Invalid input. Please enter 1 or 2. / Entrada inválida. Por favor, insira 1 ou 2.")

def list_files_and_folders(directory, mode="B", list_option=1, recursive=False, specific_subfolders=None, ignore_hidden=False):
    folder_name = os.path.basename(os.path.normpath(directory))
    disk_letter = os.path.splitdrive(os.path.abspath(directory))[0].replace(":", "")
    output_filename_base = f"{folder_name} ({disk_letter})"

    # Gather folders and files
    folders = []
    files = []
    import re
    for entry in os.scandir(directory):
        if ignore_hidden and is_hidden_file(entry):
            continue
        if entry.is_dir():
            if not specific_subfolders:
                folders.append(entry.path)
            else:
                entry_name_clean = re.sub(r'\W+', '', entry.name).lower()
                for sub in specific_subfolders:
                    sub_clean = re.sub(r'\W+', '', sub).lower()
                    if sub_clean == entry_name_clean:
                        folders.append(entry.path)
                        break
        elif entry.is_file():
            files.append(entry.path)

    def get_all_sub_files(folder):
        sub_files = []
        if not specific_subfolders:
            for root, _, file_list in os.walk(folder):
                for f in file_list:
                    sub_files.append(os.path.join(root, f))
        else:
            for entry in os.scandir(folder):
                if entry.is_file():
                    sub_files.append(entry.path)
        return sub_files

    sub_files_map = {}
    if recursive:
        for folder in folders:
            sub_files_map[folder] = get_all_sub_files(folder)

    # Output
    output_file_path = os.path.join(directory, f"{output_filename_base}.{'docx' if mode.upper() == 'A' else ('json' if mode.upper() == 'C' else 'txt')}")
    # Detect language for output credits
    # Use global LANGUAGES and lang if available, else fallback to English
    import inspect
    frame = inspect.currentframe()
    while frame:
        if 'lang' in frame.f_globals:
            lang = frame.f_globals['lang']
            break
        frame = frame.f_back
    else:
        lang = 'en'
    L = LANGUAGES[lang] if 'LANGUAGES' in globals() and lang in LANGUAGES else LANGUAGES['en']
    credits_str = L["credits"]
    if mode.upper() == "A":
        try:
            doc = Document()
            doc.add_heading(f"{folder_name}", level=1)
            # Folders
            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_docx(doc, folder, list_option=list_option)
            # Files
            if list_option == 1:
                for file in files:
                    p = doc.add_paragraph()
                    p.add_run("• ")
                    base, ext = os.path.splitext(os.path.basename(file))
                    run = p.add_run(base)
                    p.add_run(ext)
            # Add credits to DOCX file (small text)
            credits_p = doc.add_paragraph()
            credits_run = credits_p.add_run(credits_str)
            credits_run.font.size = 127000  # 8pt in EMU units (1 pt = 12700 EMU)
            doc.save(output_file_path)
            print(f"List generated successfully: {output_file_path}")
            print("The List has been generated")
        except Exception as e:
            print(f"DOCX generation failed: {e}. Falling back to TXT mode.")
            output_file_path = os.path.join(directory, f"{output_filename_base}.txt")
            with open(output_file_path, "w", encoding="utf-8") as txt_file:
                txt_file.write(f"{folder_name}\n\n")
                # Folders
                if list_option in [1, 2]:
                    for folder in folders:
                        write_folder_structure_txt(txt_file, folder)
                # Files
                if list_option in [1, 3]:
                    for file in files:
                        base, ext = os.path.splitext(os.path.basename(file))
                        txt_file.write(f"• {base}{ext}\n")
                txt_file.write(f"\n\n{credits_str}\n")
            print(f"List generated successfully: {output_file_path}")
            print("The List has been generated")
    elif mode.upper() == "C":
        # JSON export
        def folder_to_dict(folder):
            d = {"folder": os.path.basename(folder), "files": [], "subfolders": []}
            for entry in os.scandir(folder):
                if entry.is_file():
                    d["files"].append(entry.name)
                elif entry.is_dir():
                    d["subfolders"].append(folder_to_dict(entry.path))
            return d
        db = {"root": folder_name, "files": [os.path.basename(f) for f in files] if list_option in [1, 3] else [], "folders": []}
        for folder in folders:
            db["folders"].append(folder_to_dict(folder))
        with open(output_file_path, "w", encoding="utf-8") as json_file:
            json.dump(db, json_file, indent=2)
            json_file.write(f"\n/* {credits_str} */\n")
        print(f"JSON database exported: {output_file_path}")
    else:
        with open(output_file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(f"{folder_name}\n\n")
            # Folders
            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_txt(txt_file, folder)
            # Files
            if list_option in [1, 3]:
                for file in files:
                    base, ext = os.path.splitext(os.path.basename(file))
                    txt_file.write(f"• {base}{ext}\n")
            txt_file.write(f"\n\n{credits_str}\n")
        print(f"List generated successfully: {output_file_path}")
        print("The List has been generated")

def write_folder_structure_docx(doc, folder, indent=0, list_option=1):
    p = doc.add_paragraph("    " * indent)
    p.add_run("• ")
    run = p.add_run(os.path.basename(folder))
    run.bold = True
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    subfolder_count = 0
    subfile_count = 0
    import inspect
    specific_subfolders = inspect.currentframe().f_back.f_locals.get('specific_subfolders', None)
    for entry in entries:
        if entry.is_dir():
            if not specific_subfolders:
                subfolder_count += 1
                write_folder_structure_docx(doc, entry.path, indent + 1, list_option)
        elif entry.is_file() and list_option == 1:
            subfile_count += 1
            sub_p = doc.add_paragraph("    " * (indent + 1) + f"{subfile_count}. ")
            base, ext = os.path.splitext(entry.name)
            sub_run = sub_p.add_run(base)
            sub_run.italic = True
            sub_p.add_run(ext)

def write_folder_structure_txt(txt_file, folder, indent=0, list_option=1):
    txt_file.write(f"{'    ' * indent}• {os.path.basename(folder)}\n")
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    subfolder_count = 0
    subfile_count = 0
    import inspect
    specific_subfolders = inspect.currentframe().f_back.f_locals.get('specific_subfolders', None)
    for entry in entries:
        if entry.is_dir():
            if not specific_subfolders:
                subfolder_count += 1
                write_folder_structure_txt(txt_file, entry.path, indent + 1, list_option)
        elif entry.is_file() and list_option == 1:
            subfile_count += 1
            base, ext = os.path.splitext(entry.name)
            txt_file.write(f"{'    ' * (indent + 1)}{subfile_count}. {base}{ext}\n")

def is_hidden_file(entry):
    # Windows hidden files: starts with '.' or has hidden attribute, or common system files
    hidden_names = {"desktop.ini", "thumbs.db", "._.ds_store", ".ds_store", ".gitignore", ".gitkeep"}
    if entry.name.startswith('.') or entry.name.lower() in hidden_names:
        return True
    try:
        import ctypes
        attrs = ctypes.windll.kernel32.GetFileAttributesW(str(entry.path))
        if attrs != -1 and attrs & 2:  # FILE_ATTRIBUTE_HIDDEN = 0x2
            return True
    except Exception:
        pass
    return False

if __name__ == "__main__":
    lang = get_lang()
    L = LANGUAGES[lang]
    while True:
        directory_to_list = input(L["enter_directory"])
        if not os.path.isdir(directory_to_list):
            print(L["error_directory"].format(dir=directory_to_list))
        else:
            break
    recursive = True  # Always list sub-folders
    hide_hidden_input = input(L["hide_hidden"]).strip().lower()
    ignore_hidden = hide_hidden_input in (["yes", "y"] if lang == "en" else ["sim", "s"])
    filter_input = input(L["filter_input"]).strip()
    if filter_input:
        specific_subfolders = [folder.strip() for folder in filter_input.split(",") if folder.strip()]
    else:
        specific_subfolders = None
    while True:
        mode = input(L["output_mode"]).strip().lower()
        if mode in (["a", "docx", "b", "txt", "c", "json"] if lang == "en" else ["a", "docx", "b", "txt", "c", "json"]):
            if mode in ["a", "docx"]:
                mode = "A"
            elif mode in ["b", "txt"]:
                mode = "B"
            else:
                mode = "C"
            break
        else:
            print(L["invalid_mode"])
    list_option = 0
    while True:
        try:
            print(L["choose_option"])
            print(L["option_1"])
            print(L["option_2"])
            print(L["option_3"])
            choice = input(L["enter_choice"])
            list_option = int(choice)
            if list_option in [1, 2, 3]:
                break
            else:
                print(L["invalid_choice"])
        except ValueError:
            print(L["invalid_number"])
    # Patch all print statements in list_files_and_folders to use L
    def patched_list_files_and_folders(*args, **kwargs):
        # Monkey patch print inside this function
        import builtins
        orig_print = builtins.print
        def lang_print(*p_args, **p_kwargs):
            # Replace known strings
            msg = p_args[0] if p_args else ""
            if isinstance(msg, str):
                msg = msg.replace("List generated successfully:", L["list_success"].split(":")[0]+":")
                msg = msg.replace("The List has been generated", L["list_generated"])
                msg = msg.replace("DOCX generation failed:", L["docx_failed"].split(":")[0]+":")
                msg = msg.replace("JSON database exported:", L["json_exported"].split(":")[0]+":")
                msg = msg.replace("Credits: User Ium101 from GitHub and AI tools", L["credits"])
            orig_print(msg, *p_args[1:], **p_kwargs)
        builtins.print = lang_print
        try:
            list_files_and_folders(*args, **kwargs)
        finally:
            builtins.print = orig_print
    patched_list_files_and_folders(directory_to_list, mode=mode, list_option=list_option, recursive=recursive, specific_subfolders=specific_subfolders, ignore_hidden=ignore_hidden)
