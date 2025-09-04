import os
import json
from docx import Document

def list_files_and_folders(directory, mode="B", list_option=1, recursive=False, specific_subfolders=None):
    folder_name = os.path.basename(os.path.normpath(directory))
    disk_letter = os.path.splitdrive(os.path.abspath(directory))[0].replace(":", "")
    output_filename_base = f"{folder_name} ({disk_letter})"

    # Gather folders and files
    folders = []
    files = []
    import re
    for entry in os.scandir(directory):
        if entry.is_dir():
            if not specific_subfolders:
                folders.append(entry.path)
            else:
                entry_name_clean = re.sub(r'\W+', '', entry.name).lower()
                for sub in specific_subfolders:
                    sub_clean = re.sub(r'\W+', '', sub).lower()
                    if sub_clean == entry_name_clean:
                        folders.append(entry.path)
                        break
        elif entry.is_file():
            files.append(entry.path)

    def get_all_sub_files(folder):
        sub_files = []
        if not specific_subfolders:
            for root, _, file_list in os.walk(folder):
                for f in file_list:
                    sub_files.append(os.path.join(root, f))
        else:
            for entry in os.scandir(folder):
                if entry.is_file():
                    sub_files.append(entry.path)
        return sub_files

    sub_files_map = {}
    if recursive:
        for folder in folders:
            sub_files_map[folder] = get_all_sub_files(folder)

    output_file_path = os.path.join(directory, f"{output_filename_base}.{'docx' if mode.upper() == 'A' else ('json' if mode.upper() == 'C' else 'txt')}")
    if mode.upper() == "A":
        try:
            doc = Document()
            doc.add_heading(f"File List for {folder_name}", level=1)
            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_docx(doc, folder)
            if list_option in [1, 3]:
                for file in files:
                    p = doc.add_paragraph()
                    p.add_run("• ")
                    base, ext = os.path.splitext(os.path.basename(file))
                    run = p.add_run(base)
                    run.bold = True
                    p.add_run(ext)
            # Add credits to DOCX file (small text)
            credits_p = doc.add_paragraph()
            credits_run = credits_p.add_run("Credits: User Lum-10 from GitHub and AI tools")
            credits_run.font.size = 127000  # 8pt in EMU units (1 pt = 12700 EMU)
            doc.save(output_file_path)
            print(f"List generated successfully: {output_file_path}")
            print("The List has been generated")
        except Exception as e:
            print(f"DOCX generation failed: {e}. Falling back to TXT mode.")
            output_file_path = os.path.join(directory, f"{output_filename_base}.txt")
            with open(output_file_path, "w", encoding="utf-8") as txt_file:
                txt_file.write(f"File List for {folder_name}\n\n")
                if list_option in [1, 2]:
                    for folder in folders:
                        write_folder_structure_txt(txt_file, folder)
                if list_option in [1, 3]:
                    for file in files:
                        base, ext = os.path.splitext(os.path.basename(file))
                        txt_file.write(f"• {base}{ext}\n")
            print(f"List generated successfully: {output_file_path}")
            print("The List has been generated")
            with open(output_file_path, "a", encoding="utf-8") as txt_file:
                txt_file.write("\n\nCredits: User Lum-10 from GitHub and AI tools\n")
    elif mode.upper() == "C":
        # JSON export
        def folder_to_dict(folder):
            d = {"folder": os.path.basename(folder), "files": [], "subfolders": []}
            for entry in os.scandir(folder):
                if entry.is_file():
                    d["files"].append(entry.name)
                elif entry.is_dir():
                    d["subfolders"].append(folder_to_dict(entry.path))
            return d
        db = {"root": folder_name, "files": [os.path.basename(f) for f in files], "folders": []}
        for folder in folders:
            db["folders"].append(folder_to_dict(folder))
        with open(output_file_path, "w", encoding="utf-8") as json_file:
            json.dump(db, json_file, indent=2)
            # Add credits as a comment at the end (not valid JSON, but for user info)
            json_file.write("\n/* Credits: User Lum-10 from GitHub and AI tools */\n")
        print(f"JSON database exported: {output_file_path}")
    else:
        with open(output_file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(f"File List for {folder_name}\n\n")
            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_txt(txt_file, folder)
            if list_option in [1, 3]:
                for file in files:
                    base, ext = os.path.splitext(os.path.basename(file))
                    txt_file.write(f"• {base}{ext}\n")
        print(f"List generated successfully: {output_file_path}")
        print("The List has been generated")
        with open(output_file_path, "a", encoding="utf-8") as txt_file:
            txt_file.write("\n\nCredits: User Lum-10 from GitHub and AI tools\n")

def write_folder_structure_docx(doc, folder, indent=0):
    p = doc.add_paragraph("    " * indent)
    p.add_run("• ")
    run = p.add_run(os.path.basename(folder))
    run.bold = True
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    subfolder_count = 0
    subfile_count = 0
    import inspect
    specific_subfolders = inspect.currentframe().f_back.f_locals.get('specific_subfolders', None)
    for entry in entries:
        if entry.is_dir():
            if not specific_subfolders:
                subfolder_count += 1
                write_folder_structure_docx(doc, entry.path, indent + 1)
        elif entry.is_file():
            subfile_count += 1
            sub_p = doc.add_paragraph("    " * (indent + 1) + f"{subfile_count}. ")
            base, ext = os.path.splitext(entry.name)
            sub_run = sub_p.add_run(base)
            sub_run.italic = True
            sub_p.add_run(ext)

def write_folder_structure_txt(txt_file, folder, indent=0):
    txt_file.write(f"{'    ' * indent}• {os.path.basename(folder)}\n")
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    subfolder_count = 0
    subfile_count = 0
    import inspect
    specific_subfolders = inspect.currentframe().f_back.f_locals.get('specific_subfolders', None)
    for entry in entries:
        if entry.is_dir():
            if not specific_subfolders:
                subfolder_count += 1
                write_folder_structure_txt(txt_file, entry.path, indent + 1)
        elif entry.is_file():
            subfile_count += 1
            base, ext = os.path.splitext(entry.name)
            txt_file.write(f"{'    ' * (indent + 1)}{subfile_count}. {base}{ext}\n")

if __name__ == "__main__":
    directory_to_list = input("Enter the directory to list: ")
    recursive = True  # Always list sub-folders
    filter_input = input("Enter sub-folder names or keywords to filter (comma-separated, case-insensitive, substring match), or leave blank to include all: ").strip()
    if filter_input:
        specific_subfolders = [folder.strip() for folder in filter_input.split(",") if folder.strip()]
    else:
        specific_subfolders = None
    while True:
        mode = input("Do you want to generate the output as Mode A / DOCX, Mode B / TXT, or Mode C / JSON (for Database)? ").strip().lower()
        if mode in ["a", "docx", "b", "txt", "c", "json"]:
            if mode in ["a", "docx"]:
                mode = "A"
            elif mode in ["b", "txt"]:
                mode = "B"
            else:
                mode = "C"
            break
        else:
            print("Invalid input. Please enter A/a/DOCX/docx for Mode A, B/b/TXT/txt for Mode B, or C/c/JSON/json for Mode C.")
    list_option = 0
    while True:
        try:
            print("Choose listing option:")
            print("  1. List both folders and files")
            print("  2. List only folders")
            print("  3. List only files")
            choice = input("Enter your choice (1/2/3): ")
            list_option = int(choice)
            if list_option in [1, 2, 3]:
                break
            else:
                print("Invalid choice. Please enter 1, 2, or 3.")
        except ValueError:
            print("Invalid input. Please enter a number (1, 2, or 3).")
    list_files_and_folders(directory_to_list, mode=mode, list_option=list_option, recursive=recursive, specific_subfolders=specific_subfolders)
