import os
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from docx import Document
import json

def is_hidden_file(entry):
    hidden_names = {"desktop.ini", "thumbs.db", "._.ds_store", ".ds_store", ".gitignore", ".gitkeep"}
    if entry.name.startswith('.') or entry.name.lower() in hidden_names:
        return True
    try:
        import ctypes
        attrs = ctypes.windll.kernel32.GetFileAttributesW(str(entry.path))
        if attrs != -1 and attrs & 2:
            return True
    except Exception:
        pass
    return False

def list_files_and_folders(directory, mode="B", list_option=1, recursive=True, specific_subfolders=None, ignore_hidden=False):
    folder_name = os.path.basename(os.path.normpath(directory))
    disk_letter = os.path.splitdrive(os.path.abspath(directory))[0].replace(":", "")
    output_filename_base = f"{folder_name} ({disk_letter})"
    folders = []
    files = []
    import re
    for entry in os.scandir(directory):
        if ignore_hidden and is_hidden_file(entry):
            continue
        if entry.is_dir():
            if not specific_subfolders:
                folders.append(entry.path)
            else:
                entry_name_clean = re.sub(r'\W+', '', entry.name).lower()
                for sub in specific_subfolders:
                    sub_clean = re.sub(r'\W+', '', sub).lower()
                    if sub_clean == entry_name_clean:
                        folders.append(entry.path)
                        break
        elif entry.is_file():
            files.append(entry.path)
    def get_all_sub_files(folder):
        sub_files = []
        if not specific_subfolders:
            for root, _, file_list in os.walk(folder):
                for f in file_list:
                    sub_files.append(os.path.join(root, f))
        else:
            for entry in os.scandir(folder):
                if entry.is_file():
                    sub_files.append(entry.path)
        return sub_files
    sub_files_map = {}
    if recursive:
        for folder in folders:
            sub_files_map[folder] = get_all_sub_files(folder)
    output_file_path = os.path.join(directory, f"{output_filename_base}.{'docx' if mode.upper() == 'A' else ('json' if mode.upper() == 'C' else 'txt')}")
    if mode.upper() == "A":
        try:
            doc = Document()
            doc.add_heading(f"{folder_name}", level=1)
            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_docx(doc, folder, list_option=list_option)
            if list_option == 1:
                for file in files:
                    p = doc.add_paragraph()
                    p.add_run("• ")
                    base, ext = os.path.splitext(os.path.basename(file))
                    run = p.add_run(base)
                    p.add_run(ext)
            credits_p = doc.add_paragraph()
            credits_run = credits_p.add_run("Credits: User Ium101 from GitHub and AI tools")
            credits_run.font.size = 127000
            doc.save(output_file_path)
            messagebox.showinfo("Success", f"List generated successfully: {output_file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"DOCX generation failed: {e}")
    elif mode.upper() == "C":
        def folder_to_dict(folder):
            d = {"folder": os.path.basename(folder), "files": [], "subfolders": []}
            for entry in os.scandir(folder):
                if entry.is_file():
                    d["files"].append(entry.name)
                elif entry.is_dir():
                    d["subfolders"].append(folder_to_dict(entry.path))
            return d
        db = {"root": folder_name, "files": [os.path.basename(f) for f in files] if list_option in [1, 3] else [], "folders": []}
        for folder in folders:
            db["folders"].append(folder_to_dict(folder))
        with open(output_file_path, "w", encoding="utf-8") as json_file:
            json.dump(db, json_file, indent=2)
            json_file.write("\n/* Credits: User Ium101 from GitHub and AI tools */\n")
        messagebox.showinfo("Success", f"JSON database exported: {output_file_path}")
    else:
        with open(output_file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(f"{folder_name}\n\n")
            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_txt(txt_file, folder, list_option=list_option)
            if list_option == 1 or list_option == 3:
                if list_option == 1:
                    for file in files:
                        base, ext = os.path.splitext(os.path.basename(file))
                        txt_file.write(f"• {base}{ext}\n")
                elif list_option == 3:
                    for file in files:
                        base, ext = os.path.splitext(os.path.basename(file))
                        txt_file.write(f"• {base}{ext}\n")
            txt_file.write("\n\nCredits: User Ium101 from GitHub and AI tools\n")
        messagebox.showinfo("Success", f"List generated successfully: {output_file_path}")

def write_folder_structure_docx(doc, folder, indent=0, list_option=1):
    p = doc.add_paragraph("    " * indent)
    p.add_run("• ")
    run = p.add_run(os.path.basename(folder))
    run.bold = True
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    subfolder_count = 0
    subfile_count = 0
    import inspect
    specific_subfolders = inspect.currentframe().f_back.f_locals.get('specific_subfolders', None)
    for entry in entries:
        if entry.is_dir():
            if not specific_subfolders:
                subfolder_count += 1
                write_folder_structure_docx(doc, entry.path, indent + 1, list_option)
        elif entry.is_file() and list_option == 1:
            subfile_count += 1
            sub_p = doc.add_paragraph("    " * (indent + 1) + f"{subfile_count}. ")
            base, ext = os.path.splitext(entry.name)
            sub_run = sub_p.add_run(base)
            sub_run.italic = True
            sub_p.add_run(ext)

def write_folder_structure_txt(txt_file, folder, indent=0, list_option=1):
    txt_file.write(f"{'    ' * indent}• {os.path.basename(folder)}\n")
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    subfolder_count = 0
    subfile_count = 0
    import inspect
    specific_subfolders = inspect.currentframe().f_back.f_locals.get('specific_subfolders', None)
    for entry in entries:
        if entry.is_dir():
            if not specific_subfolders:
                subfolder_count += 1
                write_folder_structure_txt(txt_file, entry.path, indent + 1, list_option)
        elif entry.is_file() and list_option == 1:
            subfile_count += 1
            base, ext = os.path.splitext(entry.name)
            txt_file.write(f"{'    ' * (indent + 1)}{subfile_count}. {base}{ext}\n")

def run_gui():
    LANGUAGES = {
        "en": {
            "title": "Lister Z",
            "run": "Run Files & Folders Lister Z",
            "select_dir": "Select a folder to list.",
            "error_dir": "The directory '{directory}' does not exist. Please select a valid folder.",
            "mode": "Do you want to generate the output as DOCX (A), TXT (B), or JSON (C)?",
            "invalid_mode": "Invalid input. Please enter DOCX/A or TXT/B, or JSON/C.",
            "list_option": "Choose listing option:\n1. Both folders and files\n2. Only folders\n3. Only files",
            "filter": "Enter sub-folder names or keywords to filter (comma-separated), or leave blank to include all:",
            "hide_hidden": "Do you want to hide files such as desktop.ini?",
            "success": "List generated successfully: {path}",
            "json_success": "JSON database exported: {path}",
            "docx_error": "DOCX generation failed: {err}",
            "credits": "Credits: User Ium101 from GitHub and AI tools"
        },
        "pt": {
            "title": "Lister Z",
            "run": "Executar Listador de Pastas e Arquivos Z",
            "select_dir": "Selecione uma pasta para listar.",
            "error_dir": "O diretório '{directory}' não existe. Por favor, selecione uma pasta válida.",
            "mode": "Deseja gerar a saída como DOCX (A), TXT (B) ou JSON (C)?",
            "invalid_mode": "Entrada inválida. Por favor, insira DOCX/A, TXT/B ou JSON/C.",
            "list_option": "Escolha a opção de listagem:\n1. Pastas e arquivos\n2. Apenas pastas\n3. Apenas arquivos",
            "filter": "Digite nomes de subpastas ou palavras-chave para filtrar (separados por vírgula), ou deixe em branco para incluir todas:",
            "hide_hidden": "Deseja ocultar arquivos como desktop.ini?",
            "success": "Lista gerada com sucesso: {path}",
            "json_success": "Banco de dados JSON exportado: {path}",
            "docx_error": "Falha ao gerar DOCX: {err}",
        "credits": "Créditos: Usuário Ium101 do GitHub e Ferramentas IA"
        }
    }
    lang = ["en"]
    def set_lang(l):
        lang[0] = l
        update_ui()
    def update_ui():
        L = LANGUAGES[lang[0]]
        root.title(L["title"])
        run_btn_en.config(text=LANGUAGES["en"]["run"])
        run_btn_pt.config(text=LANGUAGES["pt"]["run"])
    root = tk.Tk()
    root.title(LANGUAGES[lang[0]]["title"])
    root.geometry("500x400")
    def select_directory():
        return filedialog.askdirectory(title=LANGUAGES[lang[0]]["select_dir"])
    def run_lister(lang_code):
        lang[0] = lang_code
        L = LANGUAGES[lang[0]]
        while True:
            directory = select_directory()
            if not directory:
                return
            if not os.path.isdir(directory):
                messagebox.showerror(L["title"], L["error_dir"].format(directory=directory))
            else:
                break
        while True:
            mode = simpledialog.askstring(L["title"], L["mode"], parent=root)
            if mode is None:
                root.destroy()
                return
            mode = mode.strip().lower()
            if mode in ["a", "docx"]:
                mode = "A"
                break
            elif mode in ["b", "txt"]:
                mode = "B"
                break
            elif mode in ["c", "json"]:
                mode = "C"
                break
            else:
                messagebox.showerror(L["title"], L["invalid_mode"])
        list_option = simpledialog.askinteger(L["title"], L["list_option"], minvalue=1, maxvalue=3, parent=root)
        if list_option is None:
            root.destroy()
            return
        filter_input = simpledialog.askstring(L["title"], L["filter"], parent=root)
        if filter_input is None:
            root.destroy()
            return
        if filter_input:
            specific_subfolders = [folder.strip() for folder in filter_input.split(",") if folder.strip()]
        else:
            specific_subfolders = None
        hide_hidden_input = messagebox.askyesno(L["title"], L["hide_hidden"], parent=root)
        if hide_hidden_input is None:
            root.destroy()
            return
        def patched_list_files_and_folders(*args, **kwargs):
            import builtins
            orig_showinfo = messagebox.showinfo
            orig_showerror = messagebox.showerror
            def lang_showinfo(title, msg, *a, **kw):
                if "List generated successfully:" in msg:
                    msg = L["success"].format(path=msg.split(": ",1)[-1])
                elif "JSON database exported:" in msg:
                    msg = L["json_success"].format(path=msg.split(": ",1)[-1])
                elif "Credits: User Ium101 from GitHub and AI tools" in msg:
                    msg = L["credits"]
                orig_showinfo(title, msg, *a, **kw)
            def lang_showerror(title, msg, *a, **kw):
                if "DOCX generation failed:" in msg:
                    msg = L["docx_error"].format(err=msg.split(": ",1)[-1])
                orig_showerror(title, msg, *a, **kw)
            messagebox.showinfo = lang_showinfo
            messagebox.showerror = lang_showerror
            try:
                list_files_and_folders(*args, **kwargs)
            finally:
                messagebox.showinfo = orig_showinfo
                messagebox.showerror = orig_showerror
        patched_list_files_and_folders(directory, mode=mode, list_option=list_option, recursive=True, specific_subfolders=specific_subfolders, ignore_hidden=hide_hidden_input)
    run_btn_pt = tk.Button(root, text=LANGUAGES["pt"]["run"], command=lambda: run_lister("pt"), height=2, width=30)
    run_btn_pt.pack(pady=(30,10))
    run_btn_en = tk.Button(root, text=LANGUAGES["en"]["run"], command=lambda: run_lister("en"), height=2, width=30)
    run_btn_en.pack(pady=(0,10))
    credits_frame = tk.Frame(root)
    credits_frame.pack(side="bottom", pady=10)
    credits_lbl_pt = tk.Label(credits_frame, text=LANGUAGES["pt"]["credits"], font=("Arial", 8))
    credits_lbl_pt.pack()
    credits_lbl_en = tk.Label(credits_frame, text=LANGUAGES["en"]["credits"], font=("Arial", 8))
    credits_lbl_en.pack()
    update_ui()
    root.mainloop()

if __name__ == "__main__":
    run_gui()
